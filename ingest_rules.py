from __future__ import annotations

import argparse
import hashlib
import json
from pathlib import Path
from typing import List
from datetime import datetime, timezone

from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma

# Config

EMBED_MODEL = "BAAI/bge-large-en-v1.5"

WANTED_DOCX = [
    "abstract_component.docx",
    "abstract_relations.docx",
    "ansible_mapping.docx",
    "kubernetes_mapping.docx",
    "terraform_mapping.docx",
    "edmm_metamodel.docx",
    "edmm_yaml_spec.docx",
]



# Utils: state + hashing


def sha256_file(path: Path, buf_size: int = 1024 * 1024) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        while True:
            b = f.read(buf_size)
            if not b:
                break
            h.update(b)
    return h.hexdigest()


def load_state(state_path: Path) -> dict:
    if not state_path.exists():
        return {}
    try:
        return json.loads(state_path.read_text(encoding="utf-8"))
    except Exception:
        return {}


def save_state(state_path: Path, state: dict) -> None:
    state["updated_at"] = datetime.now(timezone.utc).isoformat()
    state_path.write_text(json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8")


# DOCX parsing


def platform_from_filename(name: str) -> str:
    n = name.lower()
    if "edmm" in n:
        return "abstract"
    if "ansible" in n:
        return "ansible"
    if "kubernetes" in n or "kube" in n or "k8" in n:
        return "kubernetes"
    if "terraform" in n:
        return "terraform"
    if "abstract" in n:
        return "abstract"
    return "unknown"


def doc_type_from_filename(name: str) -> str:
    n = name.lower()
    if "edmm_metamodel" in n:
        return "edmm_metamodel"
    if "edmm_yaml_spec" in n:
        return "edmm_yaml_spec"
    if "abstract_component" in n:
        return "abstract_component"
    if "abstract_relation" in n or "abstract_relations" in n:
        return "abstract_relations"
    if "mapping" in n:
        return "mapping_rule"
    return "unknown"


def read_docx_lines_header_aware(docx_path: Path) -> List[str]:
    """
    Output list of strings suitable for embedding.
    - Tables: use header row to create key=value pairs per row.
    - Paragraphs: keep as one block to preserve context.
    """
    d = DocxDocument(str(docx_path))
    out: List[str] = []

    # Tables (header-aware)
    for table in d.tables:
        rows = table.rows
        if not rows:
            continue

        header = [c.text.strip() for c in rows[0].cells]
        header = [h if h else f"col{i+1}" for i, h in enumerate(header)]

        # if only header exists
        if len(rows) == 1:
            out.append("TABLE_HEADER: " + " | ".join(header))
            continue

        for r in rows[1:]:
            cells = [c.text.strip() for c in r.cells]
            if all(not c for c in cells):
                continue

            # align lengths
            if len(cells) < len(header):
                cells += [""] * (len(header) - len(cells))
            if len(cells) > len(header):
                cells = cells[:len(header)]

            pairs = []
            for h, v in zip(header, cells):
                if v:
                    pairs.append(f"{h}={v}")
            line = "; ".join(pairs).strip()
            if len(line) >= 10:
                out.append(line)

    # Paragraph block
    paras = [p.text.strip() for p in d.paragraphs if p.text and p.text.strip()]
    if paras:
        out.append("\n".join(paras))

    # uniq
    uniq, seen = [], set()
    for x in out:
        k = x.strip()
        if k and k not in seen:
            seen.add(k)
            uniq.append(k)
    return uniq



# File selection (allowlist only)

def select_docx_files(rules_dir: Path) -> List[Path]:
    files: List[Path] = []
    for fname in WANTED_DOCX:
        p = rules_dir / fname
        if p.exists() and p.is_file():
            files.append(p)

    files = sorted(set(files))
    return files



# Build collection


def build_collection(
    rules_dir: Path,
    persist_dir: Path,
    collection_name: str,
    chunk_size: int,
    chunk_overlap: int,
    reset: bool,
    dry_run: bool,
):
    files = select_docx_files(rules_dir)
    if not files:
        raise FileNotFoundError(
            f"None of the wanted DOCX files found under: {rules_dir}\nExpected: {WANTED_DOCX}"
        )

    print(f"\n[kb_core] selected files: {len(files)}")
    for f in files:
        print(" -", f)

    state_path = persist_dir / f"state_{collection_name}.json"
    state = {} if reset else load_state(state_path)
    state.setdefault("files", {})
    state.setdefault("params", {})
    state["params"] = {
        "embed_model": EMBED_MODEL,
        "chunk_size": chunk_size,
        "chunk_overlap": chunk_overlap,
    }

    if dry_run:
        print("[kb_core] dry_run=True -> not indexing.")
        return

    to_index: List[Path] = []
    skipped = 0
    for p in files:
        h = sha256_file(p)
        prev = state["files"].get(str(p))
        if prev and prev.get("sha256") == h:
            skipped += 1
            continue
        to_index.append(p)
        state["files"][str(p)] = {"sha256": h}

    print(f"[kb_core] unchanged skipped: {skipped}  to_index: {len(to_index)}")

    embeddings = HuggingFaceEmbeddings(
        model_name=EMBED_MODEL,
        encode_kwargs={"normalize_embeddings": True},
    )

    db = Chroma(
        collection_name=collection_name,
        persist_directory=str(persist_dir),
        embedding_function=embeddings,
    )

    if reset:
        try:
            db.delete_collection()
            print(f"[kb_core] ✅ deleted old collection ({collection_name})")
        except Exception:
            pass
        db = Chroma(
            collection_name=collection_name,
            persist_directory=str(persist_dir),
            embedding_function=embeddings,
        )
    else:
        # delete only changed sources
        for p in to_index:
            try:
                db.delete(where={"source": str(p)})
            except Exception:
                pass

    docs: List[Document] = []
    for p in to_index:
        lines = read_docx_lines_header_aware(p)
        platform = platform_from_filename(p.name)
        dtype = doc_type_from_filename(p.name)

        for i, line in enumerate(lines, 1):
            docs.append(
                Document(
                    page_content=line,
                    metadata={
                        "source": str(p),
                        "filename": p.name,
                        "platform": platform,
                        "doc_type": dtype,
                        "item_id": f"{p.stem}_{i}",
                    },
                )
            )

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        add_start_index=True,
    )
    chunks = splitter.split_documents(docs)
    print(f"[kb_core] docs: {len(docs)}  chunks: {len(chunks)}")

    if chunks:
        db.add_documents(chunks)

    save_state(state_path, state)
    print(f"[kb_core] ✅ saved: {persist_dir}  collection={collection_name}")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--rules_dir", default="/app/Evaluation/Evaluation", help="Directory containing DOCX files")
    ap.add_argument("--persist_dir", default="/app/project_folder/chroma_db")
    ap.add_argument("--collection", default="kb_core")
    ap.add_argument("--dry_run", action="store_true")
    ap.add_argument("--reset", action="store_true")
    ap.add_argument("--chunk_size", type=int, default=900)
    ap.add_argument("--chunk_overlap", type=int, default=90)
    args = ap.parse_args()

    rules_dir = Path(args.rules_dir)
    persist_dir = Path(args.persist_dir)
    persist_dir.mkdir(parents=True, exist_ok=True)

    build_collection(
        rules_dir=rules_dir,
        persist_dir=persist_dir,
        collection_name=args.collection,
        chunk_size=args.chunk_size,
        chunk_overlap=args.chunk_overlap,
        reset=args.reset,
        dry_run=args.dry_run,
    )


if __name__ == "__main__":
    main()
